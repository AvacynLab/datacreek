# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the terms described in the LICENSE file in
# the root directory of this source tree.
# DOCX parasers
import os
from typing import Any, Dict

from .base import BaseParser


class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents"""

    def parse(self, file_path: str, *, use_unstructured: bool = True) -> str:
        """Parse a DOCX file into plain text

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text from the document
        """
        if use_unstructured:
            try:
                from unstructured.partition.docx import partition_docx

                elements = partition_docx(filename=file_path)
                texts = [
                    getattr(el, "text", str(el)) for el in elements if getattr(el, "text", None)
                ]
                return "\n".join(texts)
            except Exception:
                pass

        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
            )

        doc = docx.Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        return "\n\n".join(p for p in paragraphs if p)

    def save(self, content: str, output_path: str) -> None:
        """Save the extracted text to a file

        Args:
            content: Extracted text content
            output_path: Path to save the text
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
